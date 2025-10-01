import os
import io
import json
import re
import math
from pathlib import Path
from collections import Counter
from datetime import datetime

import streamlit as st

# File parsing
from docx import Document as DocxDocument
import pdfplumber

# PDF builder
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4

# OpenAI (latest SDK)
from openai import OpenAI

# ---------------------------
# Constants / stopwords
# ---------------------------
STOPWORDS = set("""
a an the and or but if then with without at by for to from of on in into over under
up down across through as is are was were be been being that this those these it its it's
they them their my our your you i me we us he she his her him who whom which what where
when why how not no nor so than too very can just also only into such per via etc vs
""".split())

# Multiple template options
TEMPLATES = {
    "Classic": """# {name}
{city} · {email} · {phone} · {linkedin}

## Summary
{summary}

## Skills
{skills_section}

## Experience
{experience_section}

## Projects
{projects_section}

## Education
{education_section}
""",
    "Minimal": """# {name}
{email} | {phone} | {linkedin} | {city}

## Professional Summary
{summary}

## Key Skills
{skills_section}

## Work Experience
{experience_section}

## Academic Background
{education_section}

## Projects
{projects_section}
""",
    "Compact": """# {name} — {city}
📧 {email} | 📱 {phone} | 🔗 {linkedin}

### Profile
{summary}

### Skills
{skills_section}

### Experience
{experience_section}

### Education
{education_section}

### Projects
{projects_section}
""",
    "Modern": """# {name}
*Location:* {city} | *Email:* {email} | *Phone:* {phone} | *LinkedIn:* {linkedin}

## About Me
{summary}

## Skills Snapshot
{skills_section}

## Work History
{experience_section}

## Notable Projects
{projects_section}

## Education
{education_section}
"""
}

# ---------------------------
# Utils
# ---------------------------
def normalize(text: str) -> str:
    text = text.lower()
    text = re.sub(r"http[s]?://\S+", " ", text)
    text = re.sub(r"[^a-z0-9+.# ]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text

def tokenize(text: str):
    return [t for t in normalize(text).split() if t not in STOPWORDS]

def extract_keywords(text: str, top_k=60):
    toks = tokenize(text)
    freq = Counter(toks)
    common = [w for w, _ in freq.most_common(top_k)]
    filtered = [w for w in common if len(w) > 2 or any(c.isdigit() for c in w) or '+' in w or '#' in w or '.' in w]
    seen, out = set(), []
    for w in filtered:
        if w not in seen:
            seen.add(w); out.append(w)
    return out

def bow_vectorize(doc_tokens, vocab):
    counts = Counter(doc_tokens)
    return [counts.get(w, 0) for w in vocab]

def cosine(v1, v2):
    dot = sum(a*b for a,b in zip(v1, v2))
    n1 = math.sqrt(sum(a*a for a in v1))
    n2 = math.sqrt(sum(b*b for b in v2))
    return 0.0 if (n1 == 0 or n2 == 0) else dot/(n1*n2)

def make_vocab(*docs):
    all_tokens = []
    for d in docs:
        all_tokens.extend(tokenize(d))
    freq = Counter(all_tokens)
    return [w for w,_ in freq.most_common(2000)]

def read_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")

def read_docx(file_bytes: bytes) -> str:
    bio = io.BytesIO(file_bytes)
    doc = DocxDocument(bio)
    return "\n".join(p.text for p in doc.paragraphs)

def read_pdf(file_bytes: bytes) -> str:
    text = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text.append(page.extract_text() or "")
    return "\n".join(text)

def parse_file(uploaded_file) -> str:
    if uploaded_file is None:
        return ""
    data = uploaded_file.read()
    name = uploaded_file.name.lower()
    if name.endswith(".txt"):
        return read_txt(data)
    if name.endswith(".docx"):
        return read_docx(data)
    if name.endswith(".pdf"):
        return read_pdf(data)
    try:
        return data.decode("utf-8", errors="ignore")
    except Exception:
        return ""

def load_role_skills(skills_db: dict, target_role: str):
    role_key = None
    tr = target_role.lower()
    for k in skills_db.keys():
        if k.lower() == tr:
            role_key = k; break
    if role_key is None:
        for k in skills_db.keys():
            if tr in k.lower() or k.lower() in tr:
                role_key = k; break
    return skills_db.get(role_key, {"core": [], "nice_to_have": []}), (role_key or target_role)

def compute_role_relevance(resume_text: str, role_skills: dict):
    resume_tokens = set(tokenize(resume_text))
    skills = [s.lower() for s in role_skills.get("core", []) + role_skills.get("nice_to_have", [])]
    skills_tokens = set(normalize(s) for s in skills)
    present = sum(1 for s in skills_tokens if s in resume_tokens)
    total = max(1, len(skills_tokens))
    missing = sorted(list(skills_tokens - resume_tokens))
    return present/total, missing

def extract_header(resume_raw: str):
    get = lambda pat, default: (re.search(pat, resume_raw, flags=re.I) or [None, default])[1].strip()
    return {
        "name": get(r"(?i)name:\s*([^\n]+)", "Your Name"),
        "email": get(r"(?i)email:\s*([^\n]+)", "you@example.com"),
        "phone": get(r"(?i)phone:\s*([^\n]+)", "+91-0000000000"),
        "city": get(r"(?i)city:\s*([^\n]+)", "Your City"),
        "linkedin": get(r"(?i)linkedin:\s*([^\n]+)", "linkedin.com/in/yourprofile"),
    }

def build_md_template(header, summary, skills_section, experience_section, projects_section, education_section, template_choice="Classic"):
    template = TEMPLATES.get(template_choice, TEMPLATES["Classic"])
    return template.format(
        name=header["name"], city=header["city"], email=header["email"], phone=header["phone"],
        linkedin=header["linkedin"], summary=summary, skills_section=skills_section,
        experience_section=experience_section, projects_section=projects_section,
        education_section=education_section
    )

# ATS-safe DOCX builder
def md_to_docx(md_text: str) -> bytes:
    doc = DocxDocument()
    for line in md_text.splitlines():
        line = line.rstrip()
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ATS-safe PDF builder
def md_to_pdf(md_text: str) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    for line in md_text.splitlines():
        line = line.strip()
        if not line:
            story.append(Spacer(1, 10))
            continue
        if line.startswith("# "):
            story.append(Paragraph(f"<b>{line[2:].strip()}</b>", styles["Title"]))
        elif line.startswith("## "):
            story.append(Paragraph(f"<b>{line[3:].strip()}</b>", styles["Heading2"]))
        elif line.startswith("- "):
            story.append(Paragraph("• " + line[2:].strip(), styles["Normal"]))
        else:
            story.append(Paragraph(line, styles["Normal"]))
        story.append(Spacer(1, 6))

    doc.build(story)
    pdf_value = buffer.getvalue()
    buffer.close()
    return pdf_value

# JSON export function
def parse_markdown_to_json(md_text: str, header_info: dict, resolved_role: str, ats_score: float) -> str:
    """
    Convert markdown resume to structured JSON format with full content
    """
    # Extract name and contact from markdown if not in header_info
    lines = md_text.split('\n')
    
    # Try to extract header from first few lines
    if lines and lines[0].startswith('# '):
        name_line = lines[0][2:].strip()
        if '—' in name_line:
            header_info['name'] = name_line.split('—')[0].strip()
        else:
            header_info['name'] = name_line
    
    # Parse contact info from second line if present
    if len(lines) > 1:
        contact_line = lines[1]
        email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', contact_line)
        phone_match = re.search(r'[\+\d][\d\s\-\(\)]+', contact_line)
        linkedin_match = re.search(r'linkedin\.com/[\w\-/]+', contact_line)
        
        if email_match:
            header_info['email'] = email_match.group()
        if phone_match:
            header_info['phone'] = phone_match.group().strip()
        if linkedin_match:
            header_info['linkedin'] = linkedin_match.group()
    
    sections = {
        "personal_info": header_info,
        "target_role": resolved_role,
        "ats_score": round(ats_score, 3),
        "summary": "",
        "skills": [],
        "experience": [],
        "projects": [],
        "education": [],
        "currently_learning": []
    }
    
    current_section = None
    current_item = {}
    summary_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip header lines (first 2-3 lines)
        if i < 3 and (line.startswith('# ') or '·' in line or '|' in line or '📧' in line):
            i += 1
            continue
        
        # Detect sections
        if line.startswith('## ') or line.startswith('### ') and len(line) > 4:
            section_title = line.replace('##', '').replace('#', '').strip().lower()
            
            # Save previous item if exists
            if current_item and current_section in ["experience", "projects", "education"]:
                sections[current_section].append(current_item)
                current_item = {}
            
            # Map section titles
            if 'summary' in section_title or 'about' in section_title or 'profile' in section_title:
                current_section = "summary"
                summary_lines = []
            elif 'skill' in section_title:
                current_section = "skills"
            elif 'experience' in section_title or 'work' in section_title or 'history' in section_title:
                current_section = "experience"
            elif 'project' in section_title:
                current_section = "projects"
            elif 'education' in section_title or 'academic' in section_title:
                current_section = "education"
            elif 'learning' in section_title:
                current_section = "currently_learning"
            else:
                current_section = None
                
        elif line and current_section:
            # Process content based on section
            if current_section == "summary":
                if not line.startswith('#') and not line.startswith('-') and not line.startswith('*'):
                    summary_lines.append(line)
                    sections["summary"] = " ".join(summary_lines)
                    
            elif current_section == "skills":
                if line.startswith('- ') or line.startswith('* '):
                    skill = line[2:].strip()
                    # Remove markdown formatting
                    skill = re.sub(r'\*\*([^*]+)\*\*', r'\1', skill)
                    skill = re.sub(r'\*([^*]+)\*', r'\1', skill)
                    skill = re.sub(r'_([^_]+)_', r'\1', skill)
                    sections["skills"].append(skill)
                elif line and not line.startswith('#'):
                    # Handle skills not in bullet format
                    skills_list = [s.strip() for s in line.split(',') if s.strip()]
                    for skill in skills_list:
                        skill = re.sub(r'\*\*([^*]+)\*\*', r'\1', skill)
                        skill = re.sub(r'\*([^*]+)\*', r'\1', skill)
                        if skill and skill not in sections["skills"]:
                            sections["skills"].append(skill)
                    
            elif current_section in ["experience", "projects", "education"]:
                if line.startswith('###') or (line.startswith('**') and line.endswith('**')):
                    # Save previous item
                    if current_item:
                        sections[current_section].append(current_item)
                    # Start new item
                    title = line.replace('###', '').replace('**', '').strip()
                    current_item = {
                        "title": title,
                        "description": "",
                        "details": []
                    }
                elif line.startswith('- ') or line.startswith('* '):
                    detail = line[2:].strip()
                    # Remove markdown formatting
                    detail = re.sub(r'\*\*([^*]+)\*\*', r'\1', detail)
                    detail = re.sub(r'\*([^*]+)\*', r'\1', detail)
                    if current_item:
                        current_item["details"].append(detail)
                    else:
                        # Bullet without title - create item
                        current_item = {
                            "title": detail,
                            "description": "",
                            "details": []
                        }
                elif line and not line.startswith('#'):
                    # Additional descriptive text
                    if current_item:
                        if current_item["description"]:
                            current_item["description"] += " " + line
                        else:
                            current_item["description"] = line
                    else:
                        # Text without item - create one
                        current_item = {
                            "title": line,
                            "description": "",
                            "details": []
                        }
                        
            elif current_section == "currently_learning":
                if line.startswith('- ') or line.startswith('* '):
                    learning = line[2:].strip()
                    learning = re.sub(r'\*\*([^*]+)\*\*', r'\1', learning)
                    sections["currently_learning"].append(learning)
        
        i += 1
    
    # Save last item if exists
    if current_item and current_section in ["experience", "projects", "education"]:
        sections[current_section].append(current_item)
    
    # Clean up
    sections["summary"] = sections["summary"].strip()
    
    # Remove empty description fields
    for section in ["experience", "projects", "education"]:
        for item in sections[section]:
            if "description" in item and not item["description"]:
                del item["description"]
    
    # Add metadata
    result = {
        "resume_data": sections,
        "metadata": {
            "generated_at": datetime.now().isoformat(),
            "format_version": "1.0",
            "ats_optimized": True
        }
    }
    
    return json.dumps(result, indent=2, ensure_ascii=False)

# ---------------------------
# OpenAI helpers
# ---------------------------
def get_openai_client():
    return OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Enhanced system prompt for better ATS optimization
REWRITE_SYS_PROMPT = """You are an expert resume writer specializing in ATS optimization with a MANDATORY goal: achieve 75%+ ATS score.

CRITICAL REQUIREMENTS:
1. The resume MUST achieve 75%+ ATS compatibility score
2. Heavily incorporate ALL target skills provided by the user
3. Use exact skill keywords from the job description and skills database
4. Optimize for keyword density while maintaining readability

Rules:
- Tailor the entire resume ONLY around the target skills provided by the user
- Use EXACT keywords from target skills list - don't paraphrase or use synonyms
- Include target skills multiple times across different sections (Summary, Skills, Experience, Projects)
- Keep formatting ATS-friendly: plain headings, bullet points, no tables/columns/icons
- Do NOT fabricate jobs, education, or achievements
- Rephrase experience, projects, and summary to strongly highlight the target skills
- Always use bullet points starting with action verbs (Developed, Implemented, Optimized, Utilized)
- If the candidate lacks a skill, mention it under 'Currently Learning' or integrate learning projects
- Output sections in this order ONLY: Summary, Skills, Experience, Projects, Education, Currently Learning (if needed)
- Maximize keyword overlap with job description and target skills

KEYWORD OPTIMIZATION STRATEGY:
- Repeat important keywords 2-3 times across different sections
- Use both acronyms and full forms (ML and Machine Learning)
- Include skill variations (Python programming, Python development)
- Add relevant tools and frameworks even if not explicitly mentioned in original resume
"""

def gpt_rewrite_resume(client: OpenAI, resume_text: str, jd_text: str, target_role: str, missing_readable: list, suggested_skills: str):
    user_prompt = f"""
Target role: {target_role}

CRITICAL: The output resume MUST achieve 75%+ ATS score. Focus heavily on keyword optimization.

User Suggested Skills (MUST include these exact keywords):
{suggested_skills}

Job Description (optimize for these keywords):
\"\"\"{jd_text}\"\"\"

User Resume (raw text):
\"\"\"{resume_text}\"\"\"

HIGH PRIORITY Missing Skills (include these EXACT keywords multiple times):
{", ".join(missing_readable) if missing_readable else "None"}

INSTRUCTIONS:
1. Use EXACT keywords from the skills list - don't paraphrase
2. Include target skills in Summary, Skills section, and throughout Experience/Projects
3. Repeat important keywords 2-3 times across different sections
4. Add relevant learning projects if skills are missing from experience
5. Optimize for maximum keyword match with job description and target skills

Rewrite the resume following the rules strictly for maximum ATS score (75%+ required).
"""
    try:
        resp = client.responses.create(
            model="gpt-4o-mini",
            instructions=REWRITE_SYS_PROMPT,
            input=user_prompt,
        )
        md = resp.output_text.strip()
    except Exception:
        comp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": REWRITE_SYS_PROMPT},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.2,
        )
        md = comp.choices[0].message.content.strip()
    return md

# New function for iterative improvement
def improve_resume_iteratively(client: OpenAI, resume_md: str, jd_text: str, target_role: str, 
                              skills_db: dict, weight_sem: float, target_score: float = 0.75, 
                              max_iterations: int = 3):
    """
    Iteratively improve resume until it reaches target ATS score (75%+)
    """
    current_resume = resume_md
    iteration = 0
    
    while iteration < max_iterations:
        # Calculate current score
        scores = compute_ats_score(current_resume, jd_text, target_role, skills_db, weight_sem)
        
        if scores['ats_score'] >= target_score:
            return current_resume, scores, iteration
        
        # If score is below target, improve it
        missing_skills = scores.get('missing_skills', [])
        role_skills, _ = load_role_skills(skills_db, target_role)
        all_target_skills = role_skills.get("core", []) + role_skills.get("nice_to_have", [])
        
        improvement_prompt = f"""
URGENT: Current ATS score is {scores['ats_score']:.3f} but MUST be 75%+ (0.75+).

Current Resume:
\"\"\"{current_resume}\"\"\"

Job Description:
\"\"\"{jd_text}\"\"\"

Missing Critical Skills: {', '.join(missing_skills)}
All Target Skills: {', '.join(all_target_skills)}

MANDATORY IMPROVEMENTS NEEDED:
1. Add missing skills to Skills section using EXACT keywords
2. Integrate missing skills into existing experience bullets
3. Create relevant project examples that showcase missing skills
4. Repeat important keywords 2-3 times across sections
5. Use both technical terms and their acronyms (e.g., "Machine Learning (ML)")

Rewrite the ENTIRE resume with aggressive keyword optimization to achieve 75%+ ATS score.
Do NOT fabricate experience but DO optimize existing content for maximum keyword density.
"""

        try:
            comp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": REWRITE_SYS_PROMPT},
                    {"role": "user", "content": improvement_prompt},
                ],
                temperature=0.1,
            )
            current_resume = comp.choices[0].message.content.strip()
        except Exception as e:
            st.error(f"Error in iteration {iteration + 1}: {e}")
            break
        
        iteration += 1
    
    # Final score calculation
    final_scores = compute_ats_score(current_resume, jd_text, target_role, skills_db, weight_sem)
    return current_resume, final_scores, iteration

# ---------------------------
# ATS Score Calculator (Before/After)
# ---------------------------
def compute_ats_score(resume_text: str, jd_text: str, target_role: str, skills_db: dict, weight_sem: float):
    jd_effective = jd_text.strip() if jd_text.strip() else f"{target_role} role"

    vocab = make_vocab(resume_text, jd_effective)
    v_resume = bow_vectorize(tokenize(resume_text), vocab)
    v_jd = bow_vectorize(tokenize(jd_effective), vocab)
    semantic_similarity = cosine(v_resume, v_jd)

    role_skills, resolved_role = load_role_skills(skills_db, target_role)
    role_relevance, missing_tokens = compute_role_relevance(resume_text, role_skills)

    final_score = round(weight_sem * semantic_similarity + (1 - weight_sem) * role_relevance, 4)

    return {
        "role": resolved_role,
        "semantic_similarity": round(semantic_similarity, 3),
        "role_relevance": round(role_relevance, 3),
        "ats_score": final_score,
        "missing_skills": missing_tokens
    }

# Enhanced keyword booster function
def boost_keywords_in_resume(resume_md: str, target_skills: list, missing_skills: list):
    """
    Enhance resume with additional keyword optimization
    """
    lines = resume_md.split('\n')
    enhanced_lines = []
    
    for line in lines:
        enhanced_lines.append(line)
        
        # Add keyword-rich bullets to Skills section
        if line.strip() == "## Skills" or line.strip() == "## Key Skills":
            # Add all target skills with variations
            skills_to_add = []
            for skill in target_skills:
                if skill.lower() not in resume_md.lower():
                    skills_to_add.append(f"- {skill}")
            
            # Add missing skills
            for skill in missing_skills:
                skill_formatted = skill.replace('_', ' ').title()
                if skill_formatted.lower() not in resume_md.lower():
                    skills_to_add.append(f"- {skill_formatted}")
            
            enhanced_lines.extend(skills_to_add)
    
    return '\n'.join(enhanced_lines)

# ---------------------------
# Streamlit UI
# ---------------------------
st.set_page_config(page_title="AI Resume Builder", page_icon="🧠", layout="wide")

st.title("🧠 AI-Powered Resume Builder (ATS-Optimized)")
st.caption("Upload your resume, add JD & skills — get a tailored, ATS-friendly resume with 75%+ score guarantee.")

with st.sidebar:
    st.header("Settings")
    api_ok = bool(os.getenv("OPENAI_API_KEY"))
    st.write("🔑 OPENAI_API_KEY " + ("✅ detected" if api_ok else "❌ missing"))
    target_role = st.text_input("Target Role", value="AI Engineer")
    suggested_skills = st.text_area("Target Skills (comma separated)", placeholder="Python, Machine Learning, Deep Learning, NLP")
    weight_sem = st.slider("Weight: Semantic Similarity", 0.0, 1.0, 0.6, 0.05)
    weight_role = 1.0 - weight_sem
    st.write(f"Weight: Role Relevance = *{weight_role:.2f}*")
    template_choice = st.selectbox("Choose Resume Template", list(TEMPLATES.keys()))
    
    # New setting for target ATS score
    target_ats_score = st.slider("Target ATS Score (%)", 75, 95, 75, 5)
    target_ats_decimal = target_ats_score / 100.0

col1, col2 = st.columns(2)
with col1:
    resume_file = st.file_uploader("Upload Resume (PDF/DOCX/TXT)", type=["pdf", "docx", "txt"])
    resume_text = parse_file(resume_file) if resume_file else st.text_area("…or paste your Resume text", height=300)
with col2:
    jd_file = st.file_uploader("Upload Job Description (optional)", type=["pdf", "docx", "txt"])
    jd_text = parse_file(jd_file) if jd_file else st.text_area("…or paste JD text (optional)", height=300, placeholder="Paste JD or leave empty")

skills_json = st.text_area("Skills DB (JSON)", height=180, value=json.dumps({
    "AI Engineer": {
        "core": ["Python", "Machine Learning", "Deep Learning", "Vector Databases", "Model Evaluation", "MLOps", "Docker", "Prompt Engineering"],
        "nice_to_have": ["LangChain", "Agents", "RAG", "LLMOps", "Kubernetes", "OpenAI API", "PyTorch", "TensorFlow"]
    },
    "Product Manager": {
        "core": ["Prioritization", "Roadmapping", "User Research", "Analytics", "A/B Testing"],
        "nice_to_have": ["SQL", "Experiment Design", "Jira", "PRD Writing"]
    }
}, indent=2))

st.divider()

run = st.button("🚀 Generate High-ATS Resume (75%+ Guaranteed)")

if run:
    if not resume_text.strip():
        st.error("Please provide a resume (upload or paste).")
        st.stop()
    try:
        skills_db = json.loads(skills_json)
    except Exception as e:
        st.error(f"Invalid Skills DB JSON: {e}")
        st.stop()

    jd_effective = jd_text.strip() if jd_text.strip() else f"{target_role} role"

    # BEFORE scores
    before_scores = compute_ats_score(resume_text, jd_effective, target_role, skills_db, weight_sem)

    role_skills, resolved_role = load_role_skills(skills_db, target_role)
    missing_readable = [s.title() for s in before_scores["missing_skills"]]
    
    # Parse suggested skills
    target_skills_list = [s.strip() for s in suggested_skills.split(',') if s.strip()]

    if not os.getenv("OPENAI_API_KEY"):
        st.warning("Set OPENAI_API_KEY in your environment to enable AI rewriting.")
        rewritten_md = build_md_template(
            extract_header(resume_text),
            summary=f"{resolved_role}–oriented professional; enable OPENAI_API_KEY for AI rewrite.",
            skills_section="- *Core:* " + ", ".join(extract_keywords(resume_text, top_k=20)),
            experience_section="- Add your experience bullets here.",
            projects_section=f"- POC for {resolved_role}\n- Learning Log",
            education_section="B.Tech / B.Sc (edit with your details)",
            template_choice=template_choice
        )
        after_scores = before_scores
    else:
        client = get_openai_client()
        
        # Initial AI rewrite
        with st.spinner("🤖 AI rewriting resume for maximum ATS optimization..."):
            rewritten_md = gpt_rewrite_resume(client, resume_text, jd_effective, resolved_role, missing_readable, suggested_skills)
        
        # Iterative improvement to reach 75%+ score
        with st.spinner("🎯 Optimizing resume to achieve 75%+ ATS score..."):
            final_resume, after_scores, iterations = improve_resume_iteratively(
                client, rewritten_md, jd_effective, target_role, skills_db, 
                weight_sem, target_ats_decimal, max_iterations=3
            )
            rewritten_md = final_resume
            
            if iterations > 0:
                st.info(f"✨ Applied {iterations} optimization iterations to reach target ATS score")

    # Enhanced keyword boosting if still below target
    if after_scores['ats_score'] < target_ats_decimal:
        with st.spinner("🔧 Applying final keyword optimizations..."):
            rewritten_md = boost_keywords_in_resume(rewritten_md, target_skills_list, after_scores.get('missing_skills', []))
            # Recalculate after boosting
            after_scores = compute_ats_score(rewritten_md, jd_effective, target_role, skills_db, weight_sem)

    # Show ATS Comparison with enhanced display
    st.subheader("📊 ATS Score Comparison")
    c1, c2 = st.columns(2)
    
    with c1:
        st.metric("Before ATS Score", f"{before_scores['ats_score']:.1%}")
        st.write("**Missing Skills:**", ", ".join(before_scores["missing_skills"][:10]) + ("..." if len(before_scores["missing_skills"]) > 10 else "") or "None")
    
    with c2:
        score_delta = after_scores['ats_score'] - before_scores['ats_score']
        st.metric("After ATS Score", f"{after_scores['ats_score']:.1%}", 
                  delta=f"{score_delta:+.1%}")
        st.write("**Remaining Missing:**", ", ".join(after_scores["missing_skills"][:5]) + ("..." if len(after_scores["missing_skills"]) > 5 else "") or "None")
    
    # Success/Warning message based on target achievement
    if after_scores['ats_score'] >= target_ats_decimal:
        st.success(f"🎉 Target achieved! ATS Score: {after_scores['ats_score']:.1%} (Target: {target_ats_score}%)")
    else:
        st.warning(f"⚠️ Target not fully reached. Current: {after_scores['ats_score']:.1%} (Target: {target_ats_score}%)")
        st.info("💡 Consider adding more relevant projects or experience that include the missing skills.")

    st.subheader("📝 Tailored Resume (Markdown)")
    st.code(rewritten_md, language="markdown")

    # Enhanced details
    with st.expander("📈 Detailed ATS Analysis"):
        st.write("**Semantic Similarity:**", f"{after_scores['semantic_similarity']:.1%}")
        st.write("**Role Relevance:**", f"{after_scores['role_relevance']:.1%}")
        st.write("**Skills Coverage:**", f"{len(before_scores['missing_skills']) - len(after_scores['missing_skills'])} additional skills matched")

    # Downloads Section
    st.subheader("📥 Download Your Resume")
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    base_filename = f"resume_{resolved_role.replace(' ','_')}_{timestamp}"
    
    # Create columns for download buttons
    dl_col1, dl_col2, dl_col3, dl_col4 = st.columns(4)
    
    with dl_col1:
        md_bytes = rewritten_md.encode("utf-8")
        st.download_button(
            "⬇ Markdown", 
            data=md_bytes, 
            file_name=f"{base_filename}.md", 
            mime="text/markdown",
            use_container_width=True
        )
    
    with dl_col2:
        docx_bytes = md_to_docx(rewritten_md)
        st.download_button(
            "⬇ DOCX", 
            data=docx_bytes, 
            file_name=f"{base_filename}.docx", 
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    
    with dl_col3:
        pdf_bytes = md_to_pdf(rewritten_md)
        st.download_button(
            "⬇ PDF", 
            data=pdf_bytes, 
            file_name=f"{base_filename}.pdf", 
            mime="application/pdf",
            use_container_width=True
        )
    
    with dl_col4:
        # Generate JSON
        header_info = extract_header(resume_text)
        json_data = parse_markdown_to_json(
            rewritten_md, 
            header_info, 
            resolved_role, 
            after_scores['ats_score']
        )
        json_bytes = json_data.encode("utf-8")
        st.download_button(
            "⬇ JSON", 
            data=json_bytes, 
            file_name=f"{base_filename}.json", 
            mime="application/json",
            use_container_width=True
        )
    
    score_emoji = "🎉" if after_scores['ats_score'] >= target_ats_decimal else "📈"
    st.success(f"{score_emoji} ATS-optimized resume generated with {after_scores['ats_score']:.1%} score!")
    
    # Optional: Show JSON preview
    with st.expander("👁️ Preview JSON Structure"):
        st.json(json.loads(json_data))